"""
contract_parser.py — Document processing for Quanby Legal Contract AI Agent
Supports PDF (via PyMuPDF) and DOCX (via python-docx)
"""

import re
import io
from typing import Optional


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text_parts = []
        for page_num, page in enumerate(doc):
            text = page.get_text("text")
            if text.strip():
                text_parts.append(f"[Page {page_num + 1}]\n{text}")
        doc.close()
        return "\n\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Failed to extract PDF text: {str(e)}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        return "\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Failed to extract DOCX text: {str(e)}")


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Auto-detect file type and extract text."""
    filename_lower = filename.lower()
    if filename_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif filename_lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    elif filename_lower.endswith(".doc"):
        raise ValueError("Legacy .doc format not supported. Please convert to .docx")
    else:
        # Try to decode as plain text
        try:
            return file_bytes.decode("utf-8")
        except Exception:
            raise ValueError(f"Unsupported file format: {filename}")


def clean_text(text: str) -> str:
    """Clean and normalize extracted text."""
    # Remove excessive whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r' {2,}', ' ', text)
    # Remove null bytes
    text = text.replace('\x00', '')
    # Strip leading/trailing whitespace per line
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)
    return text.strip()


def chunk_text(text: str, max_chunk_size: int = 8000) -> list[str]:
    """
    Split text into chunks for AI processing.
    Tries to split on paragraph boundaries.
    """
    if len(text) <= max_chunk_size:
        return [text]
    
    chunks = []
    paragraphs = text.split('\n\n')
    current_chunk = []
    current_size = 0
    
    for para in paragraphs:
        para_size = len(para)
        if current_size + para_size > max_chunk_size and current_chunk:
            chunks.append('\n\n'.join(current_chunk))
            current_chunk = [para]
            current_size = para_size
        else:
            current_chunk.append(para)
            current_size += para_size + 2  # +2 for \n\n
    
    if current_chunk:
        chunks.append('\n\n'.join(current_chunk))
    
    return chunks


def get_contract_summary_for_context(text: str, max_length: int = 12000) -> str:
    """
    Return a truncated version of the contract text suitable for AI context.
    Prioritizes the beginning and end of the document (where key terms typically appear).
    """
    if len(text) <= max_length:
        return text
    
    # Take first 2/3 and last 1/3 of the allowance
    first_portion = int(max_length * 0.66)
    last_portion = max_length - first_portion
    
    first_part = text[:first_portion]
    last_part = text[-last_portion:]
    
    return (
        first_part
        + f"\n\n[... CONTRACT CONTINUES — {len(text) - max_length} characters omitted for brevity ...]\n\n"
        + last_part
    )
